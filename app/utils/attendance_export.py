import io
import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from bson import ObjectId

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_attendance_data(db, year, month):
    month_prefix = f"{year}-{month}"
    
    # Fetch all employees
    employees = list(db.employees.find({"tenant_id": "semco"}))
    emp_map = {str(e["_id"]): e for e in employees}
    
    # Fetch attendance records for the month
    records = list(db.attendance.find({"date": {"$regex": f"^{month_prefix}"}}))
    
    # Create Summary Data
    summary_list = []
    detailed_list = []
    
    for emp_id, emp in emp_map.items():
        emp_code = emp.get("employee_code", "-")
        name = emp.get("name", "Unknown")
        designation = emp.get("designation", "-")
        
        # Stats counts
        present_count = 0
        weekly_off_count = 0
        sick_leave_count = 0
        casual_leave_count = 0
        privileged_leave_count = 0
        site_visit_count = 0
        extended_work_count = 0
        
        present_times = []
        weekly_off_times = []
        sick_leave_times = []
        casual_leave_times = []
        privileged_leave_times = []
        site_visit_times = []
        extended_work_times = []
        
        # Filter records for this employee
        emp_recs = [r for r in records if str(r["employee_id"]) == emp_id]
        
        for r in emp_recs:
            date_str = r["date"]
            day = date_str.split('-')[2]
            selections = r.get("selections", {})
            
            if 'Present' in selections:
                present_count += 1
                present_times.append(f"{day}th: {selections['Present']}")
                detailed_list.append({
                    "Date": date_str,
                    "Employee Code": emp_code,
                    "Employee Name": name,
                    "Designation": designation,
                    "Status": "Present",
                    "Time": selections['Present']
                })
                
            if 'Weekly Off' in selections:
                weekly_off_count += 1
                weekly_off_times.append(f"{day}th: {selections['Weekly Off']}")
                detailed_list.append({
                    "Date": date_str,
                    "Employee Code": emp_code,
                    "Employee Name": name,
                    "Designation": designation,
                    "Status": "Weekly Off",
                    "Time": selections['Weekly Off']
                })
                
            if 'Sick Leave' in selections:
                sick_leave_count += 1
                sick_leave_times.append(f"{day}th: {selections['Sick Leave']}")
                detailed_list.append({
                    "Date": date_str,
                    "Employee Code": emp_code,
                    "Employee Name": name,
                    "Designation": designation,
                    "Status": "Sick Leave",
                    "Time": selections['Sick Leave']
                })
                
            if 'Casual Leave' in selections:
                casual_leave_count += 1
                casual_leave_times.append(f"{day}th: {selections['Casual Leave']}")
                detailed_list.append({
                    "Date": date_str,
                    "Employee Code": emp_code,
                    "Employee Name": name,
                    "Designation": designation,
                    "Status": "Casual Leave",
                    "Time": selections['Casual Leave']
                })
                
            if 'Privileged Leave' in selections:
                privileged_leave_count += 1
                privileged_leave_times.append(f"{day}th: {selections['Privileged Leave']}")
                detailed_list.append({
                    "Date": date_str,
                    "Employee Code": emp_code,
                    "Employee Name": name,
                    "Designation": designation,
                    "Status": "Privileged Leave",
                    "Time": selections['Privileged Leave']
                })
                
            if 'Site Visit' in selections or 'Back From Site Visit' in selections:
                site_visit_count += 1
                times_sub = []
                if 'Site Visit' in selections:
                    times_sub.append(f"{selections['Site Visit']} (Out)")
                    detailed_list.append({
                        "Date": date_str,
                        "Employee Code": emp_code,
                        "Employee Name": name,
                        "Designation": designation,
                        "Status": "Site Visit",
                        "Time": selections['Site Visit']
                    })
                if 'Back From Site Visit' in selections:
                    times_sub.append(f"{selections['Back From Site Visit']} (In)")
                    detailed_list.append({
                        "Date": date_str,
                        "Employee Code": emp_code,
                        "Employee Name": name,
                        "Designation": designation,
                        "Status": "Back From Site Visit",
                        "Time": selections['Back From Site Visit']
                    })
                site_visit_times.append(f"{day}th: " + ", ".join(times_sub))
                
            if 'Extended Work' in selections:
                extended_work_count += 1
                extended_work_times.append(f"{day}th: {selections['Extended Work']}")
                detailed_list.append({
                    "Date": date_str,
                    "Employee Code": emp_code,
                    "Employee Name": name,
                    "Designation": designation,
                    "Status": "Extended Work",
                    "Time": selections['Extended Work']
                })
                
        # Format strings for cells
        p_str = f"{present_count} (" + ", ".join(present_times) + ")" if present_times else "0"
        wo_str = f"{weekly_off_count} (" + ", ".join(weekly_off_times) + ")" if weekly_off_times else "0"
        sl_str = f"{sick_leave_count} (" + ", ".join(sick_leave_times) + ")" if sick_leave_times else "0"
        cl_str = f"{casual_leave_count} (" + ", ".join(casual_leave_times) + ")" if casual_leave_times else "0"
        pl_str = f"{privileged_leave_count} (" + ", ".join(privileged_leave_times) + ")" if privileged_leave_times else "0"
        sv_str = f"{site_visit_count} (" + ", ".join(site_visit_times) + ")" if site_visit_times else "0"
        ew_str = f"{extended_work_count} (" + ", ".join(extended_work_times) + ")" if extended_work_times else "0"
        
        summary_list.append({
            "Employee Code": emp_code,
            "Employee Name": name,
            "Designation": designation,
            "Present": p_str,
            "Weekly Off": wo_str,
            "Sick Leave": sl_str,
            "Casual Leave": cl_str,
            "Privileged Leave": pl_str,
            "Site Visit": sv_str,
            "Extended Work": ew_str
        })
        
    return summary_list, sorted(detailed_list, key=lambda x: (x["Date"], x["Employee Name"]))

def export_csv(summary_data, detailed_data):
    df = pd.DataFrame(detailed_data)
    if df.empty:
        df = pd.DataFrame(columns=["Date", "Employee Code", "Employee Name", "Designation", "Status", "Time"])
    return df.to_csv(index=False).encode('utf-8')

def export_xls(summary_data, detailed_data):
    summary_df = pd.DataFrame(summary_data)
    detailed_df = pd.DataFrame(detailed_data)
    
    if summary_df.empty:
        summary_df = pd.DataFrame(columns=["Employee Code", "Employee Name", "Designation", "Present", "Weekly Off", "Sick Leave", "Casual Leave", "Privileged Leave", "Site Visit", "Extended Work"])
    if detailed_df.empty:
        detailed_df = pd.DataFrame(columns=["Date", "Employee Code", "Employee Name", "Designation", "Status", "Time"])
        
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        summary_df.to_excel(writer, sheet_name='Monthly Summary', index=False)
        detailed_df.to_excel(writer, sheet_name='Detailed Log Trail', index=False)
    output.seek(0)
    return output.getvalue()

def export_word(summary_data, detailed_data, year, month):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"SEMCO Groups\nAttendance Masterlist - {month}/{year}\n")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(16, 185, 129)
    
    doc.add_paragraph(f"Generated at: {datetime.datetime.now().strftime('%Y-%m-%d %I:%M %p')}")
    doc.add_paragraph("").paragraph_format.space_after = Pt(12)
    
    h1 = doc.add_paragraph()
    r1 = h1.add_run("Employee Log / Month Summary")
    r1.font.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(51, 65, 85)
    
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ["Emp Code", "Employee Name", "Designation", "Present", "Weekly Off", "Sick Leave", "Casual Leave", "Privileged Leave", "Site Visit", "Extended Work"]
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_background(hdr_cells[i], "10B981")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    for item in summary_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get("Employee Code", ""))
        row_cells[1].text = str(item.get("Employee Name", ""))
        row_cells[2].text = str(item.get("Designation", ""))
        row_cells[3].text = str(item.get("Present", "0"))
        row_cells[4].text = str(item.get("Weekly Off", "0"))
        row_cells[5].text = str(item.get("Sick Leave", "0"))
        row_cells[6].text = str(item.get("Casual Leave", "0"))
        row_cells[7].text = str(item.get("Privileged Leave", "0"))
        row_cells[8].text = str(item.get("Site Visit", "0"))
        row_cells[9].text = str(item.get("Extended Work", "0"))
        
    doc.add_paragraph("").paragraph_format.space_after = Pt(24)
    
    h2 = doc.add_paragraph()
    r2 = h2.add_run("Detailed Check-in Log Trail")
    r2.font.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(51, 65, 85)
    
    table2 = doc.add_table(rows=1, cols=6)
    table2.style = 'Table Grid'
    
    hdr_cells2 = table2.rows[0].cells
    headers2 = ["Date", "Emp Code", "Employee Name", "Designation", "Status", "Time"]
    for i, title_text in enumerate(headers2):
        hdr_cells2[i].text = title_text
        set_cell_background(hdr_cells2[i], "3B82F6")
        hdr_cells2[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    for item in detailed_data:
        row_cells = table2.add_row().cells
        row_cells[0].text = str(item.get("Date", ""))
        row_cells[1].text = str(item.get("Employee Code", ""))
        row_cells[2].text = str(item.get("Employee Name", ""))
        row_cells[3].text = str(item.get("Designation", ""))
        row_cells[4].text = str(item.get("Status", ""))
        row_cells[5].text = str(item.get("Time", ""))
        
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def export_pdf(summary_data, detailed_data, year, month):
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    story = []
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=colors.HexColor('#10B981'),
        alignment=1,
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        textColor=colors.HexColor('#334155'),
        spaceBefore=15,
        spaceAfter=8
    )
    
    cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7,
        leading=9
    )
    
    cell_header_style = ParagraphStyle(
        'TableHeaderCell',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white
    )

    story.append(Paragraph(f"SEMCO Groups - Monthly Attendance Report ({month}/{year})", title_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Employee Log / Month Summary", subtitle_style))
    
    summary_table_data = [[
        Paragraph("Emp Code", cell_header_style),
        Paragraph("Employee Name", cell_header_style),
        Paragraph("Designation", cell_header_style),
        Paragraph("Present", cell_header_style),
        Paragraph("Weekly Off", cell_header_style),
        Paragraph("Sick Leave", cell_header_style),
        Paragraph("Casual Leave", cell_header_style),
        Paragraph("Privileged Leave", cell_header_style),
        Paragraph("Site Visit", cell_header_style),
        Paragraph("Extended Work", cell_header_style)
    ]]
    
    for item in summary_data:
        summary_table_data.append([
            Paragraph(str(item.get("Employee Code", "-")), cell_style),
            Paragraph(str(item.get("Employee Name", "-")), cell_style),
            Paragraph(str(item.get("Designation", "-")), cell_style),
            Paragraph(str(item.get("Present", "0")), cell_style),
            Paragraph(str(item.get("Weekly Off", "0")), cell_style),
            Paragraph(str(item.get("Sick Leave", "0")), cell_style),
            Paragraph(str(item.get("Casual Leave", "0")), cell_style),
            Paragraph(str(item.get("Privileged Leave", "0")), cell_style),
            Paragraph(str(item.get("Site Visit", "0")), cell_style),
            Paragraph(str(item.get("Extended Work", "0")), cell_style)
        ])
        
    t1 = Table(summary_table_data, colWidths=[55, 90, 80, 75, 70, 70, 75, 75, 75, 75])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#10B981')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('TOPPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#F8FAFC'), colors.white])
    ]))
    
    story.append(t1)
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("Detailed Check-in Log Trail", subtitle_style))
    
    detailed_table_data = [[
        Paragraph("Date", cell_header_style),
        Paragraph("Emp Code", cell_header_style),
        Paragraph("Employee Name", cell_header_style),
        Paragraph("Designation", cell_header_style),
        Paragraph("Status", cell_header_style),
        Paragraph("Time", cell_header_style)
    ]]
    
    for item in detailed_data:
        detailed_table_data.append([
            Paragraph(str(item.get("Date", "-")), cell_style),
            Paragraph(str(item.get("Employee Code", "-")), cell_style),
            Paragraph(str(item.get("Employee Name", "-")), cell_style),
            Paragraph(str(item.get("Designation", "-")), cell_style),
            Paragraph(str(item.get("Status", "-")), cell_style),
            Paragraph(str(item.get("Time", "-")), cell_style)
        ])
        
    t2 = Table(detailed_table_data, colWidths=[90, 80, 120, 120, 120, 100])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#3B82F6')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('TOPPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#F8FAFC'), colors.white])
    ]))
    
    story.append(t2)
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
