import os
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def format_currency(val):
    if val is None:
        return "0.00"
    return f"{val:,.2f}"

def get_logo_path():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    # Primary: backend/static/logo.png
    path1 = os.path.abspath(os.path.join(current_dir, "..", "..", "static", "logo.png"))
    if os.path.exists(path1):
        return path1
    # Secondary: backend/app/static/logo.png
    path2 = os.path.abspath(os.path.join(current_dir, "..", "static", "logo.png"))
    if os.path.exists(path2):
        return path2
    return None

def generate_pdf(calc_result, employee_data, location, output_path):
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = []
    
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    
    company_style = ParagraphStyle(
        'CompanyTitle',
        parent=normal_style,
        fontName='Helvetica-Bold',
        fontSize=15,
        alignment=1, # Center
        textColor=colors.HexColor("#1e293b")
    )
    
    sub_style = ParagraphStyle(
        'CompanySub',
        parent=normal_style,
        fontName='Helvetica-Bold',
        fontSize=10,
        alignment=1,
        textColor=colors.HexColor("#64748b")
    )

    logo_path = get_logo_path()
    
    # 1. Header Section
    if logo_path:
        logo_img = Image(logo_path, width=120, height=60)
        header_table = Table([[logo_img, Paragraph("SEMCORP Process & Vacuum Systems Pvt. Ltd.<br/><font size=8>PUNE, MAHARASHTRA, INDIA</font>", company_style)]], colWidths=[140, 380])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (1,0), (1,0), 'CENTER'),
        ]))
        story.append(header_table)
    else:
        story.append(Paragraph("SEMCORP Process & Vacuum Systems Pvt. Ltd.", company_style))
        
    story.append(Spacer(1, 15))
    
    # 2. Metadata Section
    date_str = datetime.date.today().strftime("%d %B %Y")
    
    meta_rows = [
        ["Date:", date_str, "Location:", location or "PUNE"]
    ]
    if employee_data:
        meta_rows.append(["Employee Name:", employee_data.get("name", "N/A"), "Employee ID:", employee_data.get("emp_id", "N/A")])
        meta_rows.append(["Designation:", employee_data.get("designation") or employee_data.get("role", "N/A"), "Department:", employee_data.get("department", "N/A")])
        
    meta_table = Table(meta_rows, colWidths=[100, 160, 100, 160])
    meta_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTNAME', (2,0), (2,-1), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#f8fafc")),
        ('BACKGROUND', (2,0), (2,-1), colors.HexColor("#f8fafc")),
        ('PADDING', (0,0), (-1,-1), 5),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    
    story.append(Paragraph("<b>SALARY STRUCTURE & CTC BREAK-UP</b>", ParagraphStyle('Sub', parent=normal_style, fontSize=10, leading=14, spaceAfter=8)))
    story.append(meta_table)
    story.append(Spacer(1, 15))
    
    # Helper to retrieve calc values
    def val_m(key):
        return format_currency(calc_result.get(key, 0))
        
    def val_y(key):
        return format_currency(calc_result.get(key, 0) * 12)

    # 3. Main Breakup Grid Table
    grid_data = [
        ["Component", "Monthly (INR)", "Yearly (INR)"],
        ["Basic Salary (50% of Gross)", val_m("basic"), val_y("basic")],
        ["HRA (40% of Basic)", val_m("hra"), val_y("hra")],
        ["Conveyance Allowance", val_m("conveyance"), val_y("conveyance")],
        ["Education Allowance", val_m("education"), val_y("education")],
        ["Medical Allowance", val_m("medical"), val_y("medical")],
        ["Special Allowance", val_m("special"), val_y("special")],
        ["Gross Salary", val_m("gross"), val_y("gross")],
        ["Statutory Bonus (8.33% of Basic)", val_m("bonus"), val_y("bonus")],
        ["Final Gross Salary", val_m("finalGross"), val_y("finalGross")],
        ["Employee Deductions", "", ""],
        ["  - PF (Employee 12%)", val_m("employeePF"), val_y("employeePF")],
        ["  - ESIC (Employee 0.75%)", val_m("employeeESIC"), val_y("employeeESIC")],
        ["  - Professional Tax (PT)", val_m("pt"), val_y("pt")],
        ["Total Deductions", val_m("totalDeductions"), val_y("totalDeductions")],
        ["Net Take Home Salary", val_m("netTakeHome"), val_y("netTakeHome")],
        ["Employer Contributions and Cost", "", ""],
        ["  - PF (Employer Contribution)", val_m("employerPF"), val_y("employerPF")],
        ["  - ESIC (Employer 3.25%)", val_m("employerESIC"), val_y("employerESIC")],
        ["  - Gratuity (4.81% of Basic)", val_m("gratuity"), val_y("gratuity")],
        ["  - Others (Insurance, Admin)", val_m("others"), val_y("others")],
        ["Total CTC of Employee", val_m("totalCTC"), val_y("totalCTC")]
    ]

    t_grid = Table(grid_data, colWidths=[240, 140, 140])
    
    # Styling list for the main grid
    style_cmds = [
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#94a3b8")),
        ('ALIGN', (0,0), (0,-1), 'LEFT'),
        ('ALIGN', (1,0), (-1,-1), 'RIGHT'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#f1f5f9")),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('PADDING', (0,0), (-1,-1), 4),
        ('FONTSIZE', (0,0), (-1,-1), 9),
    ]

    # Gross Salary Row (slightly green) - Index 7
    style_cmds.append(('BACKGROUND', (0,7), (-1,7), colors.HexColor("#e6f4ea")))
    style_cmds.append(('FONTNAME', (0,7), (0,7), 'Helvetica-Bold'))
    
    # Final Gross Row - Index 9
    style_cmds.append(('FONTNAME', (0,9), (0,9), 'Helvetica-Bold'))
    style_cmds.append(('BACKGROUND', (0,9), (-1,9), colors.HexColor("#f8fafc")))
    
    # Employee Deductions Header - Index 10
    style_cmds.append(('SPAN', (0,10), (2,10)))
    style_cmds.append(('FONTNAME', (0,10), (-1,10), 'Helvetica-Bold'))
    style_cmds.append(('TEXTCOLOR', (0,10), (-1,10), colors.HexColor("#f97316")))
    style_cmds.append(('BACKGROUND', (0,10), (-1,10), colors.HexColor("#ffedd5")))
    
    # Total Deductions - Index 14
    style_cmds.append(('FONTNAME', (0,14), (0,14), 'Helvetica-Bold'))
    style_cmds.append(('BACKGROUND', (0,14), (-1,14), colors.HexColor("#f8fafc")))
    
    # Net Take Home - Index 15
    style_cmds.append(('FONTNAME', (0,15), (-1,15), 'Helvetica-Bold'))
    style_cmds.append(('TEXTCOLOR', (1,15), (-1,15), colors.HexColor("#cca43b"))) # Cool-toned yellow ochre
    
    # Employer Contributions Header - Index 16
    style_cmds.append(('SPAN', (0,16), (2,16)))
    style_cmds.append(('FONTNAME', (0,16), (-1,16), 'Helvetica-Bold'))
    style_cmds.append(('TEXTCOLOR', (0,16), (-1,16), colors.HexColor("#3b82f6")))
    style_cmds.append(('BACKGROUND', (0,16), (-1,16), colors.HexColor("#eff6ff")))
    
    # Total CTC row - Index 21
    style_cmds.append(('FONTNAME', (0,21), (-1,21), 'Helvetica-Bold'))
    style_cmds.append(('TEXTCOLOR', (0,21), (-1,21), colors.HexColor("#10b981"))) # Highlighted green
    style_cmds.append(('BACKGROUND', (0,21), (-1,21), colors.HexColor("#f3e8ff"))) # Translucent violet
    
    t_grid.setStyle(TableStyle(style_cmds))
    story.append(t_grid)
    
    doc.build(story)


def generate_excel(calc_result, employee_data, location, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "CTC Breakup"
    ws.views.sheetView[0].showGridLines = True
    
    # 1. Header with logo
    logo_path = get_logo_path()
    if logo_path:
        img = openpyxl.drawing.image.Image(logo_path)
        img.width = 130
        img.height = 65
        ws.add_image(img, "A1")
        
    ws["B1"] = "SEMCORP Process & Vacuum Systems Pvt. Ltd."
    ws["B1"].font = Font(name="Arial", size=14, bold=True)
    ws["B2"] = "PUNE, MAHARASHTRA, INDIA"
    ws["B2"].font = Font(name="Arial", size=9, italic=True)
    
    # Spacer
    ws.row_dimensions[1].height = 25
    ws.row_dimensions[2].height = 18
    ws.row_dimensions[3].height = 10
    
    # 2. Metadata Block
    date_str = datetime.date.today().strftime("%d %B %Y")
    
    ws["A4"] = "Date:"
    ws["B4"] = date_str
    ws["C4"] = "Location:"
    ws["D4"] = location or "PUNE"
    
    if employee_data:
        ws["A5"] = "Employee Name:"
        ws["B5"] = employee_data.get("name", "N/A")
        ws["C5"] = "Employee ID:"
        ws["D5"] = employee_data.get("emp_id", "N/A")
        
        ws["A6"] = "Designation:"
        ws["B6"] = employee_data.get("designation") or employee_data.get("role", "N/A")
        ws["C6"] = "Department:"
        ws["D6"] = employee_data.get("department", "N/A")
        
    # Formatting Metadata Block
    border_side = Side(style='thin', color='CBD5E1')
    thin_border = Border(left=border_side, right=border_side, top=border_side, bottom=border_side)
    meta_font_bold = Font(name="Arial", size=9, bold=True)
    meta_font_normal = Font(name="Arial", size=9)
    fill_meta = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    
    for r in range(4, 7 if employee_data else 5):
        for col in ["A", "B", "C", "D"]:
            cell = ws[f"{col}{r}"]
            cell.border = thin_border
            if col in ["A", "C"]:
                cell.font = meta_font_bold
                cell.fill = fill_meta
            else:
                cell.font = meta_font_normal

    # 3. Main Grid
    start_row = 8
    headers = ["Component", "Monthly (INR)", "Yearly (INR)"]
    for c_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=c_idx, value=h)
        cell.font = Font(name="Arial", size=10, bold=True)
        cell.fill = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
        cell.alignment = Alignment(horizontal="left" if c_idx == 1 else "right")
        
    def val_m(key):
        return calc_result.get(key, 0.0)
        
    def val_y(key):
        return calc_result.get(key, 0.0) * 12

    rows_data = [
        ("Basic Salary (50% of Gross)", val_m("basic"), val_y("basic")),
        ("HRA (40% of Basic)", val_m("hra"), val_y("hra")),
        ("Conveyance Allowance", val_m("conveyance"), val_y("conveyance")),
        ("Education Allowance", val_m("education"), val_y("education")),
        ("Medical Allowance", val_m("medical"), val_y("medical")),
        ("Special Allowance", val_m("special"), val_y("special")),
        ("Gross Salary", val_m("gross"), val_y("gross")),
        ("Statutory Bonus (8.33% of Basic)", val_m("bonus"), val_y("bonus")),
        ("Final Gross Salary", val_m("finalGross"), val_y("finalGross")),
        ("Employee Deductions", None, None),
        ("  - PF (Employee 12%)", val_m("employeePF"), val_y("employeePF")),
        ("  - ESIC (Employee 0.75%)", val_m("employeeESIC"), val_y("employeeESIC")),
        ("  - Professional Tax (PT)", val_m("pt"), val_y("pt")),
        ("Total Deductions", val_m("totalDeductions"), val_y("totalDeductions")),
        ("Net Take Home Salary", val_m("netTakeHome"), val_y("netTakeHome")),
        ("Employer Contributions and Cost", None, None),
        ("  - PF (Employer Contribution)", val_m("employerPF"), val_y("employerPF")),
        ("  - ESIC (Employer 3.25%)", val_m("employerESIC"), val_y("employerESIC")),
        ("  - Gratuity (4.81% of Basic)", val_m("gratuity"), val_y("gratuity")),
        ("  - Others (Insurance, Admin)", val_m("others"), val_y("others")),
        ("Total CTC of Employee", val_m("totalCTC"), val_y("totalCTC")),
    ]

    grid_border = Border(
        left=Side(style='thin', color='94A3B8'),
        right=Side(style='thin', color='94A3B8'),
        top=Side(style='thin', color='94A3B8'),
        bottom=Side(style='thin', color='94A3B8')
    )

    for offset, row_tup in enumerate(rows_data, 1):
        curr_row = start_row + offset
        comp, m_val, y_val = row_tup
        
        c1 = ws.cell(row=curr_row, column=1, value=comp)
        c2 = ws.cell(row=curr_row, column=2, value=m_val)
        c3 = ws.cell(row=curr_row, column=3, value=y_val)
        
        # Format alignment & borders
        c1.alignment = Alignment(horizontal="left")
        c2.alignment = Alignment(horizontal="right")
        c3.alignment = Alignment(horizontal="right")
        c1.border = grid_border
        c2.border = grid_border
        c3.border = grid_border
        
        # Number formats
        if m_val is not None:
            c2.number_format = '#,##0.00'
            c3.number_format = '#,##0.00'
            
        # Fonts & Fills
        c1.font = Font(name="Arial", size=9)
        c2.font = Font(name="Arial", size=9)
        c3.font = Font(name="Arial", size=9)
        
        # Specific row styles
        if offset == 7: # Gross Salary
            c1.font = Font(name="Arial", size=9, bold=True)
            for c in [c1, c2, c3]:
                c.fill = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid")
        elif offset == 9: # Final Gross
            c1.font = Font(name="Arial", size=9, bold=True)
            for c in [c1, c2, c3]:
                c.fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
        elif offset == 10: # Employee Deductions Header
            ws.merge_cells(start_row=curr_row, start_column=1, end_row=curr_row, end_column=3)
            c1.font = Font(name="Arial", size=9, bold=True, color="F97316")
            for col_idx in [1, 2, 3]:
                cell = ws.cell(row=curr_row, column=col_idx)
                cell.fill = PatternFill(start_color="FFEDD5", end_color="FFEDD5", fill_type="solid")
                cell.border = grid_border
        elif offset == 14: # Total Deductions
            c1.font = Font(name="Arial", size=9, bold=True)
            for c in [c1, c2, c3]:
                c.fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
        elif offset == 15: # Net Take Home
            c1.font = Font(name="Arial", size=9, bold=True)
            c2.font = Font(name="Arial", size=9, bold=True, color="CCA43B")
            c3.font = Font(name="Arial", size=9, bold=True, color="CCA43B")
        elif offset == 16: # Employer Contributions Header
            ws.merge_cells(start_row=curr_row, start_column=1, end_row=curr_row, end_column=3)
            c1.font = Font(name="Arial", size=9, bold=True, color="3B82F6")
            for col_idx in [1, 2, 3]:
                cell = ws.cell(row=curr_row, column=col_idx)
                cell.fill = PatternFill(start_color="EFF6FF", end_color="EFF6FF", fill_type="solid")
                cell.border = grid_border
        elif offset == 21: # Total CTC
            c1.font = Font(name="Arial", size=9, bold=True, color="10B981")
            c2.font = Font(name="Arial", size=9, bold=True, color="10B981")
            c3.font = Font(name="Arial", size=9, bold=True, color="10B981")
            for c in [c1, c2, c3]:
                c.fill = PatternFill(start_color="F3E8FF", end_color="F3E8FF", fill_type="solid")

    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 18
    
    wb.save(output_path)


def generate_word(calc_result, employee_data, location, output_path):
    doc = Document()
    
    # Adjust margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # 1. Header with logo
    logo_path = get_logo_path()
    if logo_path:
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(2.0)
        header_table.columns[1].width = Inches(5.5)
        
        # Add logo
        cell_logo = header_table.cell(0, 0)
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.8))
        
        # Add text
        cell_text = header_table.cell(0, 1)
        p_text = cell_text.paragraphs[0]
        p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_text.add_run("SEMCORP Process & Vacuum Systems Pvt. Ltd.\n")
        run_title.font.name = "Arial"
        run_title.font.size = Pt(14)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(30, 41, 59)
        
        run_sub = p_text.add_run("PUNE, MAHARASHTRA, INDIA")
        run_sub.font.name = "Arial"
        run_sub.font.size = Pt(9)
        run_sub.font.italic = True
        run_sub.font.color.rgb = RGBColor(100, 116, 139)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p.add_run("SEMCORP Process & Vacuum Systems Pvt. Ltd.")
        run_title.font.name = "Arial"
        run_title.font.size = Pt(14)
        run_title.font.bold = True
        
    doc.add_paragraph() # Spacing
    
    # 2. Metadata Block
    date_str = datetime.date.today().strftime("%d %B %Y")
    
    meta_rows = [
        ("Date:", date_str, "Location:", location or "PUNE")
    ]
    if employee_data:
        meta_rows.append(("Employee Name:", employee_data.get("name", "N/A"), "Employee ID:", employee_data.get("emp_id", "N/A")))
        meta_rows.append(("Designation:", employee_data.get("designation") or employee_data.get("role", "N/A"), "Department:", employee_data.get("department", "N/A")))
        
    meta_table = doc.add_table(rows=len(meta_rows), cols=4)
    meta_table.autofit = False
    col_widths = [Inches(1.5), Inches(2.25), Inches(1.5), Inches(2.25)]
    
    for r_idx, row_tup in enumerate(meta_rows):
        row = meta_table.rows[r_idx]
        for c_idx, val in enumerate(row_tup):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.text = val
            
            # Apply styling
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            
            # Shading for label columns
            if c_idx in [0, 2]:
                run.font.bold = True
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>'
                cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
            
            # Borders
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(r'''
                <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                </w:tcBorders>
            ''')
            tcPr.append(borders)
            
    doc.add_paragraph() # Spacing
    
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("SALARY STRUCTURE & CTC BREAK-UP")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(10)
    r_sub.font.bold = True
    
    # 3. Main Grid Breakup
    def val_m(key):
        return format_currency(calc_result.get(key, 0))
        
    def val_y(key):
        return format_currency(calc_result.get(key, 0) * 12)

    grid_rows = [
        ("Component", "Monthly (INR)", "Yearly (INR)"),
        ("Basic Salary (50% of Gross)", val_m("basic"), val_y("basic")),
        ("HRA (40% of Basic)", val_m("hra"), val_y("hra")),
        ("Conveyance Allowance", val_m("conveyance"), val_y("conveyance")),
        ("Education Allowance", val_m("education"), val_y("education")),
        ("Medical Allowance", val_m("medical"), val_y("medical")),
        ("Special Allowance", val_m("special"), val_y("special")),
        ("Gross Salary", val_m("gross"), val_y("gross")),
        ("Statutory Bonus (8.33% of Basic)", val_m("bonus"), val_y("bonus")),
        ("Final Gross Salary", val_m("finalGross"), val_y("finalGross")),
        ("Employee Deductions", "", ""),
        ("  - PF (Employee 12%)", val_m("employeePF"), val_y("employeePF")),
        ("  - ESIC (Employee 0.75%)", val_m("employeeESIC"), val_y("employeeESIC")),
        ("  - Professional Tax (PT)", val_m("pt"), val_y("pt")),
        ("Total Deductions", val_m("totalDeductions"), val_y("totalDeductions")),
        ("Net Take Home Salary", val_m("netTakeHome"), val_y("netTakeHome")),
        ("Employer Contributions and Cost", "", ""),
        ("  - PF (Employer Contribution)", val_m("employerPF"), val_y("employerPF")),
        ("  - ESIC (Employer 3.25%)", val_m("employerESIC"), val_y("employerESIC")),
        ("  - Gratuity (4.81% of Basic)", val_m("gratuity"), val_y("gratuity")),
        ("  - Others (Insurance, Admin)", val_m("others"), val_y("others")),
        ("Total CTC of Employee", val_m("totalCTC"), val_y("totalCTC")),
    ]

    t_grid = doc.add_table(rows=len(grid_rows), cols=3)
    t_grid.autofit = False
    g_widths = [Inches(3.5), Inches(2.0), Inches(2.0)]
    
    for r_idx, row_tup in enumerate(grid_rows):
        row = t_grid.rows[r_idx]
        for c_idx, val in enumerate(row_tup):
            cell = row.cells[c_idx]
            cell.width = g_widths[c_idx]
            cell.text = val
            
            # Alignments
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
            run = p.runs[0] if p.runs else p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            
            # Borders
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(r'''
                <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="94A3B8"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="94A3B8"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="94A3B8"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="94A3B8"/>
                </w:tcBorders>
            ''')
            tcPr.append(borders)
            
            # Header Row
            if r_idx == 0:
                run.font.bold = True
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>'
                tcPr.append(parse_xml(shading_xml))
            
            # Specific row styles
            if r_idx == 7: # Gross Salary
                run.font.bold = True
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="E6F4EA"/>'
                tcPr.append(parse_xml(shading_xml))
            elif r_idx == 9: # Final Gross
                run.font.bold = True
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>'
                tcPr.append(parse_xml(shading_xml))
            elif r_idx == 10: # Employee Deductions Header
                run.font.bold = True
                run.font.color.rgb = RGBColor(249, 115, 22)
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="FFEDD5"/>'
                tcPr.append(parse_xml(shading_xml))
            elif r_idx == 14: # Total Deductions
                run.font.bold = True
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>'
                tcPr.append(parse_xml(shading_xml))
            elif r_idx == 15: # Net Take Home
                run.font.bold = True
                if c_idx > 0:
                    run.font.color.rgb = RGBColor(204, 164, 59) # Ochre
            elif r_idx == 16: # Employer Contributions Header
                run.font.bold = True
                run.font.color.rgb = RGBColor(59, 130, 246)
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="EFF6FF"/>'
                tcPr.append(parse_xml(shading_xml))
            elif r_idx == 21: # Total CTC
                run.font.bold = True
                run.font.color.rgb = RGBColor(16, 185, 129)
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="F3E8FF"/>'
                tcPr.append(parse_xml(shading_xml))
                
        # Merge cells for headers
        if r_idx in [10, 16]:
            row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            
    doc.save(output_path)
